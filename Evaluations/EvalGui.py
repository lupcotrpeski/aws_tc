import sys

if (sys.version_info > (3, 0)):
    import tkinter as tk
    from tkinter import filedialog
else:
    import Tkinter as tk
    import tkFileDialog as filedialog

import csv
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert




def process_file(filename):
    try:
        with open(filename, mode='r') as evalfile:
            reader = csv.reader(evalfile)
            ncols = len(next(reader)) + 1
            evalfile.seek(0)
            mapval = {'strongly agree': 5, 'agree': 4, 'neutral': 3, 'disagree': 2, 'strongly disagree': 1,
                      'extremely satisfied': 5, 'satisfied': 4, 'dissatisfied': 2, 'extremely dissatisfied': 1,
                      'very likely': 5, 'likely': 4, 'unlikely': 2, 'very unlikely': 1}
            ignore_vals = ["No thanks", "Yes, I'd like Amazon Web Services (AWS) to follow up with me", "Promoter", "Passive"]
            column_sum = [0] * ncols
            divisor = [0] * ncols
            question = [None] * ncols
            row_position = 0
            feedback = ''
            for row in reader:
                row_position = row_position + 1
                item_position = 0
                for item in row:
                    item_position = item_position + 1
                    if row_position == 2:
                        question[item_position] = item.replace("\n", " ")
                    if mapval.get(item.lower()):
                        column_sum[item_position] = column_sum[item_position] + mapval.get(item.lower())
                        divisor[item_position] = divisor[item_position] + 1
                    elif item and row_position > 2 and item not in ignore_vals:
                        feedback = feedback + ' - ' + item + '\n'

            item_position = 0

            instructor_sum = 0
            instructor_div = 0
            overall_sum = 0
            overall_div = 0

            output = 'Number of responses: ' + str(row_position - 2) + '\n\n'

            for pos in range(ncols):
                val = divisor[pos]
                if val > 0:
                    output += '%.2f' % (float(column_sum[pos]) / float(val)) + '\t' + question[pos] + '\n'
                    overall_div = overall_div + val
                    overall_sum = overall_sum + column_sum[pos]
                    if 'instructor' in question[pos]:
                        instructor_div = instructor_div + val
                        instructor_sum = instructor_sum + column_sum[pos]

            output += '\n'
            instructure_csat = '%.2f' % (float(instructor_sum) / float(instructor_div))
            overall_csat = '%.2f' % (float(overall_sum) / float(overall_div))
            output += instructure_csat + '\t' + 'Instructor CSAT' + '\n'
            output += overall_csat + '\t' + 'Overall CSAT' + '\n'
            output += '\n'
            output += 'Recommended Changes' + '\n'
            output += '-------------------' + '\n'
            output += feedback
            results.delete("1.0", tk.END)
            results.insert(tk.END, output)

            # Generate Word Document
            generate_docx(feedback, instructure_csat, overall_csat)
    except:
        results.insert(tk.END, "There is a problem with that file.")


def get_file_name():
    browse_file = filedialog.askopenfilename(title="Select file", filetypes=(
        ("csv file", "*.csv"), ("text file", "*.txt"), ("all files", "*.*")))
    process_file(browse_file)


def copy_text():
    r.clipboard_clear()
    r.clipboard_append(results.get("1.0", tk.END))

def generate_docx(feedback, instructure_csat, overall_csat):

    document = Document()

    document.add_picture('logo.png', width=Inches(1.25))
    paragraph = document.add_paragraph("Class Report", style='Heading 1')
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    table = document.add_table(rows=3, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Class Name'
    hdr_cells[1].text = ''
    hdr_cells[2].text = 'Class Location'
    hdr_cells[3].text = ''
    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = 'Class Type'
    hdr_cells[1].text = ''
    hdr_cells[2].text = 'Trainer'
    hdr_cells[3].text = ''
    hdr_cells = table.rows[2].cells
    hdr_cells[0].text = 'Start Date'
    hdr_cells[1].text = ''
    hdr_cells[2].text = 'End Date'
    hdr_cells[3].text = ''

    #document.add_paragraph( results.get("1.0", tk.END) )

    document.add_heading('Evaluation Summary:', level=2)
    table = document.add_table(rows=2, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Class Room'
    hdr_cells[1].text = ''
    hdr_cells[2].text = 'Content'
    hdr_cells[3].text = ''
    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = 'Instructor'
    hdr_cells[1].text = instructure_csat
    hdr_cells[2].text = 'Overall Satisfaction'
    hdr_cells[3].text = overall_csat

    document.add_heading('Student Feedback:', level=2)
    document.add_paragraph( feedback )

    document.save('classreport.docx')
    convert("classreport.docx", "classreport.pdf")



# set-up window
global r, output
r = tk.Tk()
r.title('AWS T&C CSAT Evaluation v1.1')
r.geometry("820x820")

fileOpenPath = tk.Button(r, text='Choose a raw evaluation file',
                         command=lambda: get_file_name())
fileOpenPath.grid(column=0, row=0, padx=5, pady=5, sticky=tk.W)

results = tk.Text(r, width=100, height=45)
results.grid(column=0, row=2, padx=5, pady=5, sticky=tk.W)
results.config(wrap=tk.WORD)


fileOpenPath = tk.Button(r, text='Output to DOCX',
                         command=lambda: generate_docx())
fileOpenPath.grid(column=0, row=0, padx=5, pady=5, sticky=tk.E)

fileOpenPath = tk.Button(r, text='Output to PDF',
                         command=lambda: convert("classreport.docx", "classreport.pdf"))
fileOpenPath.grid(column=0, row=1, padx=5, pady=5, sticky=tk.E)

copy = tk.Button(r, text='Copy all to clipboard', command=lambda: copy_text())
copy.grid(column=0, row=3, padx=5, pady=5, sticky=tk.W)

r.mainloop()
