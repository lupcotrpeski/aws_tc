from __future__ import print_function  # Python 2/3 compatibility
import json
import os
import requests
from lib.getEnvVariable import getEnvVariable
import sys
import csv
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert


class GenerateReport:

    # Global Variables
    apiBase = getEnvVariable('API_GATEWAY')
    environment = getEnvVariable('ENVIRONMENT')
    headers = {'Content-Type': 'application/json',
           'Authorization': 'None'}


    def __init__(self, filePath, fileName, className, classLocation, classType, trainer, startDate, endDate):
        self.fileName = fileName
        self.className = className
        self.trainer = trainer
        self.classLocation = classLocation
        self.classType = classType
        self.startDate = startDate
        self.endDate = endDate
        self.filePath = filePath


    def generateReport(self):
        self.process_file()
        return "Fudged"

    def process_file(self):
        print("Process FileName = ",self.fileName)
        try:
            csvFile = self.filePath + self.fileName
            docxFile = self.filePath + os.path.splitext(self.fileName)[0] + ".docx"
            pdfFile = self.filePath + os.path.splitext(self.fileName)[0] + ".pdf"
            print("CSV: ", csvFile)
            print("PDF: ", pdfFile)
            print("DOCX: ", docxFile)

            with open(csvFile, mode='r') as evalfile:
                reader = csv.reader(evalfile)
                ncols = len(next(reader)) + 1
                evalfile.seek(0)
                mapval = {'strongly agree': 5, 'agree': 4, 'neutral': 3, 'disagree': 2, 'strongly disagree': 1,
                        'extremely satisfied': 5, 'satisfied': 4, 'dissatisfied': 2, 'extremely dissatisfied': 1,
                        'very likely': 5, 'likely': 4, 'unlikely': 2, 'very unlikely': 1}
                ignore_vals = ["No thanks", "Yes, I'd like Amazon Web Services (AWS) to follow up with me", "Promoter", "Passive"]
                column_sum = [0] * ncols
                divisor = [0] * ncols
                question = [None] * ncols
                row_position = 0
                feedback = ''
                for row in reader:
                    row_position = row_position + 1
                    item_position = 0
                    for item in row:
                        item_position = item_position + 1
                        if row_position == 2:
                            question[item_position] = item.replace("\n", " ")
                        if mapval.get(item.lower()):
                            column_sum[item_position] = column_sum[item_position] + mapval.get(item.lower())
                            divisor[item_position] = divisor[item_position] + 1
                        elif item and row_position > 2 and item not in ignore_vals:
                            feedback = feedback + ' - ' + item + '\n'

                item_position = 0

                instructor_sum = 0
                instructor_div = 0
                content_sum = 0
                content_div = 0
                overall_sum = 0
                overall_div = 0

                output = 'Number of responses: ' + str(row_position - 2) + '\n\n'

                for pos in range(ncols):
                    val = divisor[pos]
                    if val > 0:
                        output += '%.2f' % (float(column_sum[pos]) / float(val)) + '\t' + question[pos] + '\n'
                        overall_div = overall_div + val
                        overall_sum = overall_sum + column_sum[pos]
                        if 'instructor' in question[pos]:
                            instructor_div = instructor_div + val
                            instructor_sum = instructor_sum + column_sum[pos]
                        if 'content' in question[pos]:
                            content_div = content_div + val
                            content_sum = content_sum + column_sum[pos]

                output += '\n'
                instructure_csat = '%.2f' % (float(instructor_sum) / float(instructor_div))
                overall_csat = '%.2f' % (float(overall_sum) / float(overall_div))
                output += instructure_csat + '\t' + 'Instructor CSAT' + '\n'
                output += overall_csat + '\t' + 'Overall CSAT' + '\n'
                output += '%.2f' % (float(instructor_sum) / float(instructor_div)) + '\t' + 'Instructor CSAT' + '\n'
                output += '%.2f' % (float(content_sum) / float(content_div)) + '\t' + 'Content CSAT' + '\n'
                output += '%.2f' % (float(overall_sum) / float(overall_div)) + '\t' + 'Overall CSAT' + '\n'
                output += '\n'
                output += 'Recommended Changes' + '\n'
                output += '-------------------' + '\n'
                output += feedback
                print( output)

                # Generate Word Document
                print("Generate the Word Document")
                self.generate_docx(docxFile, pdfFile, feedback, instructure_csat, overall_csat)
        except:
            print( "There is a problem with that file.")

    def generate_docx(self, docxFile, pdfFile, feedback, instructure_csat, overall_csat):
        print("PDF: ", pdfFile)
        print("DOCX: ", docxFile)
        document = Document()

        #document.add_picture('logo.png', width=Inches(1.25))
        paragraph = document.add_paragraph("Class Report", style='Heading 1')
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        table = document.add_table(rows=3, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Class Name'
        hdr_cells[1].text = ''
        hdr_cells[2].text = 'Class Location'
        hdr_cells[3].text = ''
        hdr_cells = table.rows[1].cells
        hdr_cells[0].text = 'Class Type'
        hdr_cells[1].text = ''
        hdr_cells[2].text = 'Trainer'
        hdr_cells[3].text = ''
        hdr_cells = table.rows[2].cells
        hdr_cells[0].text = 'Start Date'
        hdr_cells[1].text = ''
        hdr_cells[2].text = 'End Date'
        hdr_cells[3].text = ''

        # document.add_paragraph( results.get("1.0", tk.END) )

        document.add_heading('Evaluation Summary:', level=2)
        table = document.add_table(rows=2, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Class Room'
        hdr_cells[1].text = ''
        hdr_cells[2].text = 'Content'
        hdr_cells[3].text = ''
        hdr_cells = table.rows[1].cells
        hdr_cells[0].text = 'Instructor'
        hdr_cells[1].text = instructure_csat
        hdr_cells[2].text = 'Overall Satisfaction'
        hdr_cells[3].text = overall_csat

        document.add_heading('Student Feedback:', level=2)
        document.add_paragraph(feedback)

        document.save(docxFile)
        convert(docxFile, pdfFile)

